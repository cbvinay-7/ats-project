import os
import re
import json
import torch
import textract
import pandas as pd
import streamlit as st
from torch import combinations
from datasets import Dataset, load_dataset
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.feature_extraction.text import TfidfVectorizer
from sentence_transformers import SentenceTransformer, util
from nltk.stem import WordNetLemmatizer
import yake
from collections import defaultdict
import spacy
import altair as alt
from textstat import flesch_reading_ease, smog_index
import docx
from pypdf import PdfReader
from streamlit_ace import st_ace
from huggingface_hub import snapshot_download  # Added for downloading models

# --- MODEL & DATASET CONFIG ---
# Define local paths for downloaded models
MODELS_DIR = "./models"
JOBBERT_PATH = os.path.join(MODELS_DIR, "TechWolf-JobBERT-v3")
MINILM_PATH = os.path.join(MODELS_DIR, "all-MiniLM-L6-v2")
CVPARSER_PATH = os.path.join(MODELS_DIR, "cv-resume-parser")
FLAN_T5_PATH = os.path.join(MODELS_DIR, "flan-t5-base")
FINE_TUNED_MODEL_PATH = './fine_tuned_hybrid_model'  # Path for the user-trained model

# List of models to download from Hugging Face
MODELS_TO_DOWNLOAD = {
    "TechWolf/JobBERT-v3": JOBBERT_PATH,
    "sentence-transformers/all-MiniLM-L6-v2": MINILM_PATH,
    "nhanv/cv_parser": CVPARSER_PATH,  # Corrected model
    "google/flan-t5-base": FLAN_T5_PATH
}

# List of datasets to download
DATASETS_TO_DOWNLOAD = {

    "jacob-hugging-face/job-descriptions": "job_descriptions.csv"
}

# --- FEATURE: RESUME VERSIONING - Define path and create directory ---
RESUME_VERSIONS_PATH = "./resume_versions"
os.makedirs(RESUME_VERSIONS_PATH, exist_ok=True)

# Conditional import for pyresparser and pdfkit
try:
    from pyresparser import ResumeParser

    PYRESPARSER_AVAILABLE = True
except ImportError:
    PYRESPARSER_AVAILABLE = False


    class ResumeParser:
        def __init__(self, *args, **kwargs):
            pass

        @staticmethod
        def get_extracted_data():
            return {}


    st.warning("`pyresparser` not found. Advanced resume parsing will use a fallback.")

from transformers import AutoTokenizer, AutoModelForMaskedLM, TrainingArguments, Trainer, \
    DataCollatorForLanguageModeling, AutoModelForSeq2SeqLM

# Ensure domain_keywords.py exists
try:
    from domain_keywords import TECH_KEYWORDS, TECH_KEYWORDS_EXPANDED, KEYWORD_CATEGORIES, BASE_STOP_KEYWORDS_FILTER
except ImportError:
    st.warning("`domain_keywords.py` not found. Using dummy keyword definitions. "
               "Create this file with your actual keywords.")
    TECH_KEYWORDS = ["python", "java", "react", "aws", "sql"]
    TECH_KEYWORDS_EXPANDED = {
        "python": ["python", "django", "flask", "numpy", "pandas"],
        "java": ["java", "spring", "hibernate", "jvm"],
        "react": ["react", "redux", "next.js", "frontend"],
        "aws": ["aws", "s3", "ec2", "lambda", "cloud"],
        "sql": ["sql", "mysql", "postgresql", "database"]
    }
    KEYWORD_CATEGORIES = {
        "Programming Languages": ["python", "java", "javascript", "c++", "go"],
        "Cloud Platforms": ["aws", "azure", "gcp", "cloud"],
        "Databases": ["sql", "mysql", "postgresql", "mongodb", "cassandra"],
        "Frontend": ["react", "angular", "vue", "html", "css"],
        "Backend": ["node.js", "django", "flask", "spring", "go"],
        "DevOps": ["docker", "kubernetes", "jenkins", "ci/cd"]
    }
    BASE_STOP_KEYWORDS_FILTER = {
        "experience", "skills", "responsibilities", "requirements", "duties included",
        "responsible for", "a", "an", "the", "and", "in", "on", "with"
    }

STRONG_ACTION_VERBS = [
    'achieved', 'accelerated', 'administered', 'advised', 'advocated', 'allocated', 'analyzed', 'approved',
    'arbitrated', 'arranged', 'architected', 'assessed', 'attained', 'audited', 'authored', 'automated',
    'balanced', 'boosted', 'budgeted', 'built', 'calculated', 'centralized', 'chaired', 'championed',
    'clarified', 'coached', 'collaborated', 'conceived', 'conceptualized', 'conducted', 'consolidated',
    'constructed', 'consulted', 'contracted', 'controlled', 'converted', 'coordinated', 'counseled',
    'created', 'cultivated', 'debugged', 'decreased', 'defined', 'delegated', 'delivered', 'demonstrated',
    'designed', 'determined', 'developed', 'devised', 'directed', 'discovered', 'documented', 'doubled',
    'drafted', 'drove', 'edited', 'eliminated', 'enabled', 'enforced', 'engineered', 'enhanced',
    'ensured', 'established', 'estimated', 'evaluated', 'executed', 'expanded', 'expedited',
    'facilitated', 'finalized', 'financed', 'formalized', 'formed', 'formulated', 'fostered', 'founded',
    'generated', 'governed', 'guided', 'halved', 'headed', 'identified', 'implemented', 'improved',
    'incorporated', 'increased', 'influenced', 'initiated', 'innovated', 'inspected', 'inspired',
    'installed', 'instituted', 'instructed', 'integrated', 'interpreted', 'interviewed', 'introduced',
    'invented', 'invested', 'investigated', 'judged', 'launched', 'led', 'lectured', 'liquidated',
    'lobbied', 'maintained', 'managed', 'manufactured', 'marketed', 'mastered', 'maximized', 'mediated',
    'mentored', 'minimized', 'modernized', 'molded', 'monitored', 'motivated', 'negotiated', 'operated',
    'orchestrated', 'ordered', 'organized', 'originated', 'overhauled', 'oversaw', 'perfected',
    'performed', 'pioneered', 'planned', 'prepared', 'presented', 'presided', 'prioritized',
    'produced', 'programmed', 'projected', 'promoted', 'proposed', 'proved', 'provided', 'published',
    'purchased', 'quantified', 'raised', 'ranked', 'rated', 'rebuilt', 'received', 'recommended',
    'reconciled', 'recruited', 'redesigned', 'reduced', 're-engineered', 'regulated', 'rehabilitated',
    'reinforced', 'remodeled', 'rendered', 'reorganized', 'repaired', 'replaced', 'reported',
    'represented', 'researched', 'resolved', 'restored', 'restructured', 'retrieved', 'revamped',
    'reviewed', 'revised', 'saved', 'scheduled', 'secured', 'selected', 'served', 'serviced', 'shaped',
    'simplified', 'solidified', 'solved', 'spearheaded', 'specified', 'spoke', 'sponsored',
    'staffed', 'standardized', 'steered', 'streamlined', 'strengthened', 'structured', 'studied',
    'succeeded', 'suggested', 'summarized', 'supervised', 'supported', 'surpassed', 'surveyed',
    'synthesized', 'systematized', 'tabulated', 'targeted', 'taught', 'terminated', 'tested',
    'trained', 'transformed', 'translated', 'tripled', 'unified', 'united', 'unraveled', 'updated',
    'upgraded', 'validated', 'verbalized', 'verified', 'visualized', 'won', 'wrote'
]
WEAK_PHRASES = [
    'responsible for', 'duties included', 'worked on', 'assisted with', 'helped with',
    'was part of', 'in charge of', 'tasked with'
]
QUANTIFICATION_PATTERNS = re.compile(r'(\d+%|\$\d+|\d+\s*million|\d+\s*thousand|\d{3,})')


# --- NEW: Automated Setup Function ---
def setup_app():
    """
    Downloads required models and datasets from Hugging Face on first run.
    """
    st.info("Performing first-time setup... This may take a few minutes.")
    os.makedirs(MODELS_DIR, exist_ok=True)
    os.makedirs(DATASET_PATH, exist_ok=True)

    # --- Download Models ---
    with st.spinner("Downloading required AI models..."):
        for model_repo, local_path in MODELS_TO_DOWNLOAD.items():
            if not os.path.exists(local_path):
                try:
                    st.write(f"Downloading {model_repo}...")
                    snapshot_download(repo_id=model_repo, local_dir=local_path)
                    st.write(f"✅ {model_repo} downloaded successfully.")
                except Exception as e:
                    st.error(f"Failed to download model {model_repo}: {e}")
        st.success("All models are ready.")

    # --- Download Datasets ---
    with st.spinner("Downloading required datasets..."):
        for dataset_repo, local_filename in DATASETS_TO_DOWNLOAD.items():
            local_path = os.path.join(DATASET_PATH, local_filename)
            if not os.path.exists(local_path):
                try:
                    st.write(f"Downloading {dataset_repo}...")
                    dataset = load_dataset(dataset_repo, split='train')
                    if local_filename.endswith(".csv"):
                        dataset.to_csv(local_path)
                    else:  # Assume jsonl
                        dataset.to_json(local_path)
                    st.write(f"✅ {dataset_repo} downloaded successfully.")
                except Exception as e:
                    st.error(f"Failed to download dataset {dataset_repo}: {e}")
        st.success("All datasets are ready.")

    st.balloons()
    st.success("First-time setup complete! The app is ready.")
    st.session_state.app_setup_done = True
    st.rerun()


@st.cache_resource
def load_generative_model():
    # Use the locally downloaded model
    tokenizer = AutoTokenizer.from_pretrained(FLAN_T5_PATH)
    model = AutoModelForSeq2SeqLM.from_pretrained(FLAN_T5_PATH)
    return tokenizer, model


@st.cache_resource
def load_spacy_model():
    return spacy.load("en_core_web_sm")


nlp = load_spacy_model()

# --- Config ---
# MODEL_PATH is now dynamic based on user selection
DATASET_PATH = "./datasets"
CHECKPOINT = "distilroberta-base"  # Fallback/Training checkpoint
DEVICE = "cuda" if torch.cuda.is_available() else "cpu"
lemmatizer = WordNetLemmatizer()
YAKE_CONFIG = {"lan": "en", "n": 3, "dedupLim": 0.9, "top": 20}

SECTION_WEIGHTS = {
    "skills": 0.3, "experience": 0.5, "education": 0.1, "projects": 0.1,
    "summary": 0.05, "certifications": 0.05
}


@st.cache_resource
def load_embedder(model_path):
    """Loads a SentenceTransformer model from a given path."""
    if os.path.isdir(model_path):
        return SentenceTransformer(model_path, device=DEVICE)
    # This fallback should ideally not be hit if setup is successful
    st.warning(f"Could not find local model at {model_path}. Using checkpoint.")
    return SentenceTransformer(CHECKPOINT, device=DEVICE)


# (All other helper functions remain the same)
def clean_text(text, exclusions=None):
    if not isinstance(text, str): return ""
    text = re.sub(r'<.*?>', '', text)
    text = re.sub(r'http\S+|www\.\S+', '', text)
    text = re.sub(r'[^\w\s.,!?"]', '', text)
    text = re.sub(r'\s+', ' ', text).strip()
    stop_words_to_remove = BASE_STOP_KEYWORDS_FILTER.copy()
    if exclusions:
        stop_words_to_remove.update(exclusions)
    for word in stop_words_to_remove:
        text = re.sub(rf'\b{re.escape(word)}\b', '', text, flags=re.IGNORECASE)
    text = re.sub(r'\s+', ' ', text).strip()
    return text.strip()


def fallback_load_training_data():
    examples = []
    skipped_rows, loaded_rows = 0, 0
    preview_samples = []

    if not os.path.exists(DATASET_PATH):
        st.warning(f"Dataset path '{DATASET_PATH}' does not exist. Cannot load training data.")
        return None

    st.markdown("### 📚 Configure Training Data Loading")
    st.info("Select columns to combine or use pairwise combinations for each file.")

    files_in_dataset_path = [f for f in os.listdir(DATASET_PATH) if f.endswith(('.csv', '.json', '.jsonl'))]
    if not files_in_dataset_path:
        st.warning(f"No training files (.csv, .json, .jsonl) found in '{DATASET_PATH}'.")
        return None

    for fname in files_in_dataset_path:
        fpath = os.path.join(DATASET_PATH, fname)
        try:
            if fname.endswith(".csv"):
                df = pd.read_csv(fpath)
                st.markdown(f"#### 📄 File: `{fname}`")
                columns = df.columns.tolist()
                if len(columns) >= 2:
                    selected_columns = st.multiselect(f"Select columns to combine for `{fname}`:", columns,
                                                      default=columns[:2], key=f"multi_{fname}")
                    use_combinations = st.checkbox(f"🔁 Use all pairwise combinations for `{fname}`", value=False,
                                                   key=f"combo_{fname}")
                    if use_combinations:
                        for i, j in combinations(selected_columns, 2):
                            for _, row in df.iterrows():
                                text1, text2 = str(row[i]), str(row[j])
                                if pd.isna(text1) or pd.isna(text2):
                                    skipped_rows += 1
                                    continue
                                combined = clean_text(text1 + " " + text2)
                                examples.append({"text": combined})
                                if len(preview_samples) < 5:
                                    preview_samples.append(combined)
                                loaded_rows += 1
                    else:
                        for _, row in df.iterrows():
                            parts = [str(row[col]) for col in selected_columns if pd.notna(row[col])]
                            if not parts:
                                skipped_rows += 1
                                continue
                            combined = clean_text(" ".join(parts))
                            examples.append({"text": combined})
                            if len(preview_samples) < 5:
                                preview_samples.append(combined)
                            loaded_rows += 1
            elif fname.endswith(".json") or fname.endswith(".jsonl"):
                with open(fpath, 'r', encoding='utf-8') as f_in:
                    lines = f_in.readlines()
                json_records = [json.loads(line.strip()) for line in lines if line.strip()]
                if not json_records:
                    continue
                keys = list(json_records[0].keys())
                st.markdown(f"#### 📄 File: `{fname}`")
                selected_keys = st.multiselect(f"Select keys to combine for `{fname}`:", keys, default=keys[:2],
                                               key=f"multi_json_{fname}")
                use_combinations = st.checkbox(f"🔁 Use all pairwise combinations for `{fname}`", value=False,
                                               key=f"combo_json_{fname}")
                if use_combinations:
                    for record in json_records:
                        for k1, k2 in combinations(selected_keys, 2):
                            text1 = str(record.get(k1, ""))
                            text2 = str(record.get(k2, ""))
                            if not text1.strip() or not text2.strip():
                                skipped_rows += 1
                                continue
                            combined = clean_text(text1 + " " + text2)
                            examples.append({"text": combined})
                            if len(preview_samples) < 5:
                                preview_samples.append(combined)
                            loaded_rows += 1
                else:
                    for record in json_records:
                        parts = [str(record.get(k, "")) for k in selected_keys if str(record.get(k, "")).strip()]
                        if not parts:
                            skipped_rows += 1
                            continue
                        combined = clean_text(" ".join(parts))
                        examples.append({"text": combined})
                        if len(preview_samples) < 5:
                            preview_samples.append(combined)
                        loaded_rows += 1
        except Exception as e:
            st.warning(f"⚠️ Could not read {fname}: {e}")

    if examples:
        st.info(f"✅ Loaded {loaded_rows} training samples. Skipped {skipped_rows} invalid rows.")
        st.markdown("### 🔍 Sample Loaded Rows")
        for i, sample in enumerate(preview_samples):
            st.code(sample, language="text")
        return Dataset.from_list(examples)
    else:
        return None


def highlight_relevant_project(resume_sections, jd_text, model):
    if 'projects' not in resume_sections or not resume_sections['projects'].strip(): return None, None
    projects_text = resume_sections['projects']
    individual_projects = [p.strip() for p in re.split(r'\n\s*\n', projects_text) if p.strip()]
    if not individual_projects: return None, None
    best_project, highest_score = "", -1
    for project in individual_projects:
        score = semantic_match(project, jd_text, model)
        if score > highest_score:
            highest_score, best_project = score, project
    return best_project, highest_score


def extract_resume_skills(resume_text, keyword_categories):
    all_skills_flat = set(skill for category_skills in keyword_categories.values() for skill in category_skills)
    found_skills = {skill for skill in all_skills_flat if
                    re.search(rf'\b{re.escape(skill)}\b', resume_text, re.IGNORECASE)}
    categorized_found_skills = defaultdict(list)
    for category, category_skills in keyword_categories.items():
        for skill in found_skills:
            if skill in category_skills:
                categorized_found_skills[category].append(skill)
    return categorized_found_skills


def generate_cover_letter(resume_text, jd_text, matched_keywords, tokenizer, model):
    focus_keywords = ", ".join(matched_keywords[:5])
    prompt = f"""As an expert career coach, write a professional and enthusiastic cover letter based on the provided resume and job description.
The cover letter should be three paragraphs long.
- Paragraph 1: Introduce the candidate, express excitement for the role.
- Paragraph 2: Connect experience to job requirements, highlighting skills like {focus_keywords}.
- Paragraph 3: Reiterate interest and include a call to action.
Create a compelling narrative.

JOB DESCRIPTION (summary):
{jd_text[:1000]}

RESUME (summary):
{resume_text[:1000]}

Generated Cover Letter:"""
    try:
        inputs = tokenizer(prompt, return_tensors="pt", max_length=1024, truncation=True)
        outputs = model.generate(**inputs, max_length=512, num_beams=5, early_stopping=True)
        return tokenizer.decode(outputs[0], skip_special_tokens=True)
    except Exception as e:
        return f"Could not generate cover letter: {e}"


def analyze_cover_letter(cover_letter_text, jd_keywords):
    feedback = []
    found_keywords = [kw for kw in jd_keywords if re.search(rf'\b{re.escape(kw)}\b', cover_letter_text, re.IGNORECASE)]
    keyword_score = round(len(found_keywords) / max(1, len(jd_keywords)) * 100)
    if keyword_score >= 75:
        feedback.append(
            f"✅ **Excellent Keyword Alignment ({keyword_score}%):** Your cover letter effectively mentions key skills from the job description.")
    elif keyword_score >= 50:
        feedback.append(
            f"⚠️ **Good Keyword Alignment ({keyword_score}%):** Consider adding a few more relevant keywords.")
    else:
        feedback.append(
            f"❌ **Low Keyword Alignment ({keyword_score}%):** Your cover letter is missing many important keywords.")
    weak_phrases_found = [phrase for phrase in WEAK_PHRASES if phrase in cover_letter_text.lower()]
    if weak_phrases_found:
        feedback.append(
            f"⚠️ **Passive Tone Detected:** Replace weak phrases like '{', '.join(weak_phrases_found)}' with stronger language.")
    else:
        feedback.append("✅ **Confident Tone:** The cover letter avoids common weak phrases.")
    return feedback


def perform_sanity_checks(resume_text):
    warnings = []
    if not re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b', resume_text): warnings.append(
        "⚠️ **No email address found.**")
    if not re.search(r'(\(?\d{3}\)?[\s.-]?)?\d{3}[\s.-]?\d{4}', resume_text): warnings.append(
        "⚠️ **No phone number found.**")
    return warnings


def analyze_bullet_points(experience_text):
    if not experience_text or not isinstance(experience_text, str): return {"quantification_score": 0,
                                                                            "action_verb_score": 0,
                                                                            "weak_bullet_points": [],
                                                                            "bullet_point_count": 0}
    bullet_points = [line.strip() for line in experience_text.split('\n') if line.strip() and len(line.strip()) > 10]
    total_bullets = len(bullet_points)
    if total_bullets == 0: return {"quantification_score": 0, "action_verb_score": 0, "weak_bullet_points": [],
                                   "bullet_point_count": 0}
    quantified_count, action_verb_count, weak_points = 0, 0, []
    for point in bullet_points:
        if QUANTIFICATION_PATTERNS.search(point): quantified_count += 1
        if any(point.lower().startswith(phrase) for phrase in WEAK_PHRASES):
            weak_points.append(point)
            continue
        doc = nlp(point)
        if (len(doc) > 0 and doc[0].lemma_.lower() in STRONG_ACTION_VERBS) or (
                len(doc) > 1 and doc[1].lemma_.lower() in STRONG_ACTION_VERBS):
            action_verb_count += 1
    quant_score = round((quantified_count / total_bullets) * 100)
    action_score = round((action_verb_count / total_bullets) * 100)
    return {"quantification_score": quant_score, "action_verb_score": action_score, "weak_bullet_points": weak_points,
            "bullet_point_count": total_bullets}


def rewrite_bullet_with_ai(bullet_point, missing_keywords_list, tokenizer, model):
    # This function is now a self-contained UI component
    st.markdown(f"**Original:** {bullet_point}")
    keyword_to_include = st.selectbox("Select a missing keyword to include (optional):",
                                      options=["None"] + missing_keywords_list,
                                      key=f"keyword_select_{bullet_point[:20]}")

    if st.button("✨ Rewrite with AI", key=f"rewrite_button_{bullet_point[:20]}"):
        with st.spinner("AI is thinking..."):
            prompt = f"""Rewrite the following resume bullet point to be more impactful. Start with a strong action verb. If a keyword is provided, integrate it. Focus on achievement.
Keyword: "{keyword_to_include if keyword_to_include != 'None' else ''}"
Original Bullet: "{bullet_point}"
Rewritten Bullet:"""
            try:
                inputs = tokenizer(prompt, return_tensors="pt")
                outputs = model.generate(**inputs, max_length=80, num_beams=5, early_stopping=True)
                rewritten_text = tokenizer.decode(outputs[0], skip_special_tokens=True)
                st.success(f"**Suggestion:** {rewritten_text}")
            except Exception as e:
                st.error(f"Could not rewrite bullet: {e}")
    st.markdown("---")


def section_parse_resume(text):
    sections, current_section = defaultdict(str), "other"
    section_keywords = {"experience": ["experience", "work history"], "education": ["education"], "skills": ["skills"],
                        "projects": ["projects"], "certifications": ["certifications"],
                        "summary": ["summary", "objective"]}
    for line in text.splitlines():
        line_clean = line.strip().lower()
        matched = False
        for sec, keywords in section_keywords.items():
            if any(k in line_clean for k in keywords):
                current_section, matched = sec, True
                break
        if not matched and line.strip(): sections[current_section] += line.strip() + " "
    return sections


def enhanced_section_parser(text):
    if PYRESPARSER_AVAILABLE:
        try:
            import tempfile
            with tempfile.NamedTemporaryFile(mode='w+', delete=False, encoding='utf-8') as temp_f:
                temp_f.write(text)
                temp_path = temp_f.name
            data = ResumeParser(temp_path).get_extracted_data()
            os.remove(temp_path)
            parsed_sections = defaultdict(str)
            if data:
                parsed_sections["skills"] = ", ".join(data.get("skills", []))
                experience_text = [
                    f"{exp.get('title', '')} at {exp.get('company', '')}: {exp.get('description', '')}" if isinstance(
                        exp, dict) else str(exp) for exp in data.get("experience", [])]
                parsed_sections["experience"] = "\n".join(experience_text)
                parsed_sections["education"] = "\n".join(data.get("education", []))
            return parsed_sections
        except Exception as e:
            return section_parse_resume(text)
    else:
        return section_parse_resume(text)


@st.cache_data(ttl=3600)
def extract_keywords(text, top_n=30):
    """Extracts top_n keywords using YAKE and expands them."""
    custom_yake = {"lan": "en", "n": 3, "dedupLim": 0.9, "top": top_n}
    extractor = yake.KeywordExtractor(**custom_yake)
    keywords = [kw.lower() for kw, _ in extractor.extract_keywords(text)]
    return list(set(keywords))  # Return a unique list


def get_active_keywords(selection, user_defined_keywords):
    keyword_sets = {"Default": TECH_KEYWORDS,
                    "Extended Tech": list({kw for sublist in TECH_KEYWORDS_EXPANDED.values() for kw in sublist}),
                    "Custom": list({kw for sublist in user_defined_keywords.values() for kw in sublist})}
    return keyword_sets.get(selection, TECH_KEYWORDS)


def tfidf_match(text1, text2):
    vectors = TfidfVectorizer().fit_transform([text1, text2])
    return round(cosine_similarity(vectors[0:1], vectors[1:2]).flatten()[0] * 100, 2)


def semantic_match(text1, text2, model):
    if not text1.strip() or not text2.strip(): return 0.0
    emb1, emb2 = model.encode(text1, convert_to_tensor=True), model.encode(text2, convert_to_tensor=True)
    cos_sim = util.cos_sim(emb1, emb2).item()
    return round((cos_sim + 1) / 2 * 100, 2)


def categorize_keywords(keywords, user_defined_keywords):
    categorized, combined_categories = {}, KEYWORD_CATEGORIES.copy()
    combined_categories.update(user_defined_keywords)
    for category, kw_list in combined_categories.items():
        matched = [kw for kw in keywords if kw in set(kw_list)]
        if matched: categorized[category] = matched
    return categorized


def analyze(resume_text, jd_text, model, user_keywords=None, weights=None, exclusions=None,
            selected_keyword_system="Default"):
    sanity_warnings = perform_sanity_checks(resume_text)
    resume_clean, jd_clean = clean_text(resume_text.lower(), exclusions), clean_text(jd_text.lower(), exclusions)
    readability_score = flesch_reading_ease(resume_text)
    readability_interpretation = "Easy to read 👍" if readability_score > 70 else "Standard text ✅" if readability_score > 60 else "Hard to read ⚠️"
    resume_sections = enhanced_section_parser(resume_text)
    bullet_analysis = analyze_bullet_points(resume_sections.get("experience", ""))
    tfidf_score = tfidf_match(resume_clean, jd_clean)
    general_jd_keywords = extract_keywords(jd_clean)
    active_jd_keyword_list = get_active_keywords(selected_keyword_system, user_keywords)
    # Filter extracted keywords against the active keyword system for relevance
    jd_keywords = [kw for kw in general_jd_keywords if kw in active_jd_keyword_list] or general_jd_keywords[:25]

    weighted_semantic_sum, weighted_keyword_sum, total_section_weight_actual = 0, 0, 0
    section_keyword_coverage_details, section_semantic_score_details = {}, {}
    for section, weight in SECTION_WEIGHTS.items():
        section_content = clean_text(resume_sections.get(section, "").lower(), exclusions)
        if section_content:
            section_sem_score = semantic_match(section_content, jd_clean, model)
            section_kw_coverage = round(
                len([kw for kw in jd_keywords if kw in section_content]) / max(1, len(jd_keywords)) * 100, 2)
            weighted_semantic_sum += section_sem_score * weight
            weighted_keyword_sum += section_kw_coverage * weight
            total_section_weight_actual += weight
            section_semantic_score_details[section], section_keyword_coverage_details[
                section] = section_sem_score, section_kw_coverage
    if total_section_weight_actual > 0:
        semantic_score, keyword_coverage = round(weighted_semantic_sum / total_section_weight_actual, 2), round(
            weighted_keyword_sum / total_section_weight_actual, 2)
    else:
        semantic_score = semantic_match(resume_clean, jd_clean, model)
        keyword_coverage = round(len([kw for kw in jd_keywords if kw in resume_clean]) / max(1, len(jd_keywords)) * 100,
                                 2)
    matched, missing = sorted([kw for kw in jd_keywords if kw in resume_clean]), sorted(
        list(set(jd_keywords) - set([kw for kw in jd_keywords if kw in resume_clean])))
    tfidf_w, semantic_w, keyword_w = weights or (30, 40, 30)
    final_score = round(
        tfidf_score * tfidf_w / 100 + semantic_score * semantic_w / 100 + keyword_coverage * keyword_w / 100, 2)
    return {"tfidf_score": tfidf_score, "semantic_score": semantic_score, "keyword_coverage": keyword_coverage,
            "final_score": final_score, "matched": matched, "missing": missing, "resume_sections": resume_sections,
            "keyword_categories": categorize_keywords(jd_keywords, user_keywords),
            "section_keyword_coverage_details": section_keyword_coverage_details,
            "section_semantic_score_details": section_semantic_score_details, "bullet_analysis": bullet_analysis,
            "readability": {"score": readability_score, "interpretation": readability_interpretation},
            "sanity_warnings": sanity_warnings}


def suggest_keyword_placement(resume_text, missing_keyword, model):
    sentences = re.split(r'(?<!\w\.\w.)(?<![A-Z][a-z]\.)(?<=\.|\?)\s', resume_text)
    if not sentences: return None
    keyword_embedding, sentence_embeddings = model.encode(missing_keyword, convert_to_tensor=True), model.encode(
        sentences, convert_to_tensor=True)
    return sentences[torch.argmax(util.cos_sim(keyword_embedding, sentence_embeddings)).item()]


def generate_interview_questions(resume_text, jd_text, matched_keywords, tokenizer, model):
    focus_keywords = ", ".join(matched_keywords[:5])
    prompt = f"""Based on the resume and job description, generate 3 technical and 2 behavioral interview questions. Focus on skills like: {focus_keywords}. Format as a markdown list.
JOB DESCRIPTION: {jd_text[:1000]}
RESUME: {resume_text[:1000]}
Generated Questions:"""
    try:
        inputs = tokenizer(prompt, return_tensors="pt", max_length=1024, truncation=True)
        outputs = model.generate(**inputs, max_length=256, num_beams=5, early_stopping=True)
        return tokenizer.decode(outputs[0], skip_special_tokens=True)
    except Exception as e:
        return f"Could not generate questions: {e}"


def check_resume_formatting(file_obj):
    warnings = []
    file_name = file_obj.name
    file_obj.seek(0)
    try:
        if file_name.endswith('.docx'):
            doc = docx.Document(file_obj)
            if doc.tables: warnings.append("📄 Use of Tables: May confuse older ATS.")
            for section in doc.sections:
                if any(p.text.strip() for p in section.header.paragraphs) or any(
                        p.text.strip() for p in section.footer.paragraphs):
                    warnings.append("📄 Text in Header/Footer: Some ATS may ignore this.")
                    break
            if doc.inline_shapes: warnings.append("🖼️ Images Detected: Invisible to ATS.")
        elif file_name.endswith('.pdf'):
            reader = PdfReader(file_obj)
            if any(p.images for p in reader.pages): warnings.append("🖼️ Images Detected: Invisible to ATS.")
            if reader.get_form_text_fields(): warnings.append(
                "📄 Interactive Elements: Can interfere with text extraction.")
    except Exception as e:
        warnings.append(f"Error processing file: {e}")
    return warnings


def parse_user_keyword_input(text):
    categories = defaultdict(set)
    try:
        for block in text.strip().split("\n\n"):
            lines = block.strip().split("\n")
            if len(lines) > 1:
                category, keywords = lines[0].strip(), {kw.strip().lower() for kw in lines[1:] if kw.strip()}
                if category and keywords: categories[category] = keywords
    except Exception as e:
        st.error(f"Error parsing custom keywords: {e}")
    return dict(categories)


# --- NEW HELPER FUNCTIONS FOR IMPROVEMENTS ---
def analyze_job_description(jd_text):
    """Provides insights and potential red flags about the job description itself."""
    insights = []
    # 1. Readability
    readability_score = smog_index(jd_text)
    if readability_score > 16:
        insights.append(
            f"⚠️ **Complex Language**: This job description is difficult to read (SMOG Index: {readability_score:.1f}). This may indicate a lack of clarity in the role's definition.")
    else:
        insights.append(
            f"✅ **Clear Language**: The job description is reasonably easy to understand (SMOG Index: {readability_score:.1f}).")

    # 2. Keyword Count
    keywords = extract_keywords(jd_text, top_n=50)
    if len(keywords) > 35:
        insights.append(
            f"⚠️ **Excessive Requirements?**: The JD lists a high number of distinct skills ({len(keywords)}). This might be a 'unicorn' role asking for too many qualifications.")

    # 3. Vagueness Check
    VAGUE_JARGON = {'synergy', 'rockstar', 'ninja', 'go-getter', 'disrupt', 'paradigm shift', 'results-driven'}
    found_jargon = {jargon for jargon in VAGUE_JARGON if re.search(rf'\b{jargon}\b', jd_text, re.IGNORECASE)}
    if len(found_jargon) > 1:
        insights.append(
            f"⚠️ **Vague Jargon**: Contains corporate jargon like `{', '.join(found_jargon)}`. Ensure you understand the concrete responsibilities behind these terms.")

    return insights


def generate_score_explanation(result, weights):
    """Creates a human-readable explanation for the final score."""
    explanation = []
    tfidf_w, semantic_w, keyword_w = weights

    # Core score components
    explanation.append(
        f"- **Semantic Match ({result['semantic_score']:.1f}%)** contributed **{result['semantic_score'] * semantic_w / 100:.1f}** points to your score. This measures how well the context of your resume aligns with the job.")
    explanation.append(
        f"- **Keyword Coverage ({result['keyword_coverage']:.1f}%)** contributed **{result['keyword_coverage'] * keyword_w / 100:.1f}** points. This is based on the presence of key terms from the JD.")
    explanation.append(
        f"- **TF-IDF Relevance ({result['tfidf_score']:.1f}%)** contributed **{result['tfidf_score'] * tfidf_w / 100:.1f}** points. This reflects the statistical relevance of matching keywords.")

    # Qualitative adjustments
    bp_analysis = result['bullet_analysis']
    if bp_analysis['action_verb_score'] > 75:
        explanation.append(
            "- 👍 Your strong use of action verbs in the experience section makes your achievements clear and impactful.")
    elif bp_analysis['action_verb_score'] < 50:
        explanation.append(
            "- ⚠️ Consider starting more bullet points with strong action verbs to better showcase your accomplishments.")

    if bp_analysis['quantification_score'] > 40:
        explanation.append(
            "- 👍 Great job quantifying your results! Using numbers and metrics makes your impact tangible.")
    elif bp_analysis['quantification_score'] < 20 and bp_analysis['bullet_point_count'] > 0:
        explanation.append(
            "- ⚠️ Try to quantify more of your achievements (e.g., using $, %, or numbers) to demonstrate your value.")

    return explanation


def main():
    st.set_page_config(page_title="Hybrid Resume Analyzer", layout="wide")

    # --- NEW: Run setup on first launch ---
    if 'app_setup_done' not in st.session_state:
        setup_app()

    st.title("🧠 Resume & JD Analyzer — Hybrid + ATS Logic")

    # --- Initialize session state ---
    if 'jd_text_content' not in st.session_state: st.session_state.jd_text_content = ""
    if 'resume_text_content' not in st.session_state: st.session_state.resume_text_content = ""
    if 'analysis_result' not in st.session_state: st.session_state.analysis_result = None
    if 'last_analyzed_resume' not in st.session_state: st.session_state.last_analyzed_resume = ""
    if 'last_analyzed_jd' not in st.session_state: st.session_state.last_analyzed_jd = ""

    def process_jd_upload():
        if st.session_state.jd_uploader is not None:
            try:
                uploaded_file = st.session_state.jd_uploader
                st.session_state.jd_text_content = textract.process(uploaded_file.read(), encoding='utf-8').decode(
                    'utf-8')
            except Exception as e:
                st.error(f"Error processing {uploaded_file.name}: {e}")

    def process_resume_upload():
        if st.session_state.resume_uploader is not None:
            try:
                uploaded_file = st.session_state.resume_uploader
                st.session_state.resume_text_content = textract.process(uploaded_file.read(), encoding='utf-8').decode(
                    'utf-8')
            except Exception as e:
                st.error(f"Error processing {uploaded_file.name}: {e}")

    tab_list = ["📄 Analyze Resume", "📊 My Skills", "✉️ Cover Letter Helper", "🧠 Train Model", "📂 Batch Compare",
                "🛠️ ATS Format Check"]
    tab1, tab5, tab6, tab2, tab3, tab4 = st.tabs(tab_list)

    with st.sidebar:
        st.header("⚙️ Analysis Configuration")

        # --- MODIFIED: Model Selection ---
        st.markdown("### 🤖 Model Selection")
        model_options = {
            "JobBERT (Specialized for HR)": JOBBERT_PATH,
            "MiniLM (Fast & General Purpose)": MINILM_PATH,
            "CV-Resume-Parser (Specific Parser)": CVPARSER_PATH
        }

        # --- IMPLEMENTATION GAP FIXED ---
        # Check if a fine-tuned model exists and add it to the options
        if os.path.isdir(FINE_TUNED_MODEL_PATH):
            model_options["⭐ My Custom Trained Model"] = FINE_TUNED_MODEL_PATH
        # --- END OF FIX ---

        selected_model_name = st.selectbox(
            "Choose an analysis model:",
            options=list(model_options.keys()),
            key="model_selector"
        )
        st.session_state.selected_model_path = model_options[selected_model_name]

        st.markdown("### 🗂️ User-defined Keyword Groups")
        keyword_input = st.text_area("Define keyword groups:", height=150, key="custom_keywords_input")
        user_defined_keywords = parse_user_keyword_input(keyword_input)

        st.markdown("### 🔑 Keyword System Selection")
        selected_keyword_system = st.radio("Choose keyword set:", ("Default", "Extended Tech", "Custom"),
                                           key="keyword_system_radio")

        st.markdown("### ⚖️ Score Weighting")
        tfidf_weight = st.slider("TF-IDF Weight", 0, 100, 30, key="tfidf_weight_slider")
        semantic_weight = st.slider("Semantic Weight", 0, 100, 40, key="semantic_weight_slider")
        keyword_weight = st.slider("Keyword Match Weight", 0, 100, 30, key="keyword_weight_slider")
        if tfidf_weight + semantic_weight + keyword_weight != 100:
            st.warning("⚠️ Total weight must equal 100.")

        st.markdown("### 🚫 Exclude Keywords")
        exclusion_input = st.text_area("List exclusion keywords:", height=100, key="exclusion_keywords_input")
        exclusions = [w.strip().lower() for w in exclusion_input.splitlines() if w.strip()]

        st.markdown("### 💾 Resume Versions")
        try:
            saved_resumes = [""] + sorted(os.listdir(RESUME_VERSIONS_PATH))
            if len(saved_resumes) > 1:
                selected_version = st.selectbox("Load a saved resume version:", options=saved_resumes)
                if st.button("Load Selected Version"):
                    if selected_version:
                        with open(os.path.join(RESUME_VERSIONS_PATH, selected_version), 'r', encoding='utf-8') as f:
                            st.session_state.resume_text_content = f.read()
                        st.success(f"Loaded '{selected_version}'")
                        st.rerun()
            else:
                st.info("No saved versions found.")
        except Exception as e:
            st.error(f"Could not load resume versions: {e}")

    with tab1:
        def run_live_analysis():
            resume_changed = st.session_state.resume_text_content != st.session_state.last_analyzed_resume
            jd_changed = st.session_state.jd_text_content != st.session_state.last_analyzed_jd

            if (
                    resume_changed or jd_changed) and st.session_state.jd_text_content and st.session_state.resume_text_content:
                with st.spinner("Analyzing..."):
                    model = load_embedder(st.session_state.selected_model_path)
                    analysis_data = analyze(
                        st.session_state.resume_text_content, st.session_state.jd_text_content, model,
                        user_keywords=user_defined_keywords, weights=(tfidf_weight, semantic_weight, keyword_weight),
                        exclusions=exclusions, selected_keyword_system=selected_keyword_system
                    )
                    if analysis_data:
                        project, score = highlight_relevant_project(analysis_data['resume_sections'],
                                                                    st.session_state.jd_text_content, model)
                        if project and score: analysis_data['highlighted_project'] = {'project': project,
                                                                                      'score': score}

                        gen_tokenizer, gen_model = load_generative_model()
                        if analysis_data.get("matched"):
                            questions = generate_interview_questions(
                                st.session_state.resume_text_content, st.session_state.jd_text_content,
                                analysis_data["matched"], gen_tokenizer, gen_model
                            )
                            analysis_data["interview_questions"] = questions

                        # Store models and results in session state for interactive features
                        st.session_state.interactive_model = model
                        st.session_state.interactive_gen_tokenizer = gen_tokenizer
                        st.session_state.interactive_gen_model = gen_model

                    st.session_state.analysis_result = analysis_data
                    st.session_state.last_analyzed_resume = st.session_state.resume_text_content
                    st.session_state.last_analyzed_jd = st.session_state.jd_text_content
            # Clear results if inputs are empty
            elif not st.session_state.jd_text_content or not st.session_state.resume_text_content:
                st.session_state.analysis_result = None

        col1, col2 = st.columns(2)
        with col1:
            st.subheader("📝 Job Description")
            st.file_uploader("Upload JD", type=["pdf", "docx", "txt"], key="jd_uploader", on_change=process_jd_upload)
            st.text_area("JD Text", height=400, key="jd_text_content")

            # --- NEW: JD INSIGHTS ---
            if st.session_state.jd_text_content:
                with st.expander("🔍 Job Description Insights", expanded=False):
                    jd_insights = analyze_job_description(st.session_state.jd_text_content)
                    for insight in jd_insights:
                        st.markdown(insight)

        with col2:
            st.subheader("📄 Interactive Resume Editor")
            st.file_uploader("Upload Resume", type=["pdf", "docx", "txt"], key="resume_uploader",
                             on_change=process_resume_upload)
            st_ace(language='text', theme='chrome', keybinding='vscode', font_size=14, height=400,
                   key="resume_text_content", auto_update=True)
            st.markdown("---")
            st.subheader("💾 Save Current Resume")
            save_col1, save_col2 = st.columns([3, 1])
            with save_col1:
                file_name = st.text_input("Filename:", placeholder="e.g., resume_for_google.txt")
            with save_col2:
                st.write(""), st.write("")
                if st.button("Save Version"):
                    if file_name and st.session_state.resume_text_content:
                        if not file_name.endswith('.txt'): file_name += '.txt'
                        try:
                            save_path = os.path.join(RESUME_VERSIONS_PATH, file_name)
                            with open(save_path, 'w', encoding='utf-8') as f:
                                f.write(st.session_state.resume_text_content)
                            st.success(f"Saved to '{file_name}'")
                        except Exception as e:
                            st.error(f"Error saving file: {e}")
                    else:
                        st.warning("Please provide a filename or content.")

        # --- IMPROVEMENT: Live analysis is triggered on content change ---
        run_live_analysis()
        st.markdown("<hr>", unsafe_allow_html=True)

        if st.session_state.analysis_result:
            result = st.session_state.analysis_result
            st.success("Analysis Complete!")
            if result.get("sanity_warnings"):
                st.subheader("🚨 Pre-flight Check")
                for warning in result["sanity_warnings"]: st.warning(warning)
                st.markdown("---")

            score_col1, score_col2, score_col3, score_col4 = st.columns(4)
            with score_col1:
                st.metric(label="✅ Final Match Score", value=f"{result['final_score']}%")
            with score_col2:
                st.metric(label="🧩 Keyword Coverage", value=f"{result['keyword_coverage']}%")
            with score_col3:
                st.metric(label=" NLP Semantic Score", value=f"{result['semantic_score']}%")
            with score_col4:
                st.metric(label="Readability (Flesch)", value=f"{result['readability']['score']:.0f}",
                          help=result['readability']['interpretation'])

            st.progress(int(result['final_score']))

            # --- NEW: SCORE EXPLANATION ---
            with st.expander("💡 Score Breakdown", expanded=False):
                explanation = generate_score_explanation(result, (tfidf_weight, semantic_weight, keyword_weight))
                for point in explanation:
                    st.markdown(point)

            kw_col1, kw_col2 = st.columns(2)
            with kw_col1:
                with st.expander("✅ Matched Keywords", expanded=True):
                    st.write(f"`{', '.join(result['matched'])}`")
            with kw_col2:
                with st.expander("❌ Missing Keywords & Suggestions", expanded=True):
                    if not result['missing']:
                        st.success("No missing keywords found!")
                    for kw in result['missing']:
                        b_col1, b_col2 = st.columns([4, 1])
                        with b_col1:
                            st.markdown(f"- `{kw}`")
                        with b_col2:
                            if st.button("Suggest", key=f"suggest_{kw}", help=f"Suggest placement for '{kw}'"):
                                with st.spinner("..."):
                                    suggestion = suggest_keyword_placement(st.session_state.resume_text_content, kw,
                                                                           st.session_state.interactive_model)
                                    if suggestion:
                                        st.info(f"**Consider adding to a sentence like this:**\n\n> '{suggestion}'")
                                    else:
                                        st.warning("Could not find suitable placement.")

            with st.expander("📝 Bullet Point Analysis", expanded=False):
                bp_analysis = result['bullet_analysis']
                st.subheader(f"Found {bp_analysis['bullet_point_count']} Bullet Points in Experience Section")
                bp_col1, bp_col2 = st.columns(2)
                bp_col1.metric("Quantification Score", f"{bp_analysis['quantification_score']}%",
                               help="Percentage of bullets with numbers, $, or %.")
                bp_col2.metric("Action Verb Score", f"{bp_analysis['action_verb_score']}%",
                               help="Percentage of bullets starting with a strong action verb.")
                # --- NEW: INTEGRATED BULLET REWRITER ---
                if bp_analysis['weak_bullet_points']:
                    st.warning("Weak Bullet Points Found (Rewrite Suggestions):")
                    for point in bp_analysis['weak_bullet_points']:
                        rewrite_bullet_with_ai(
                            point, result['missing'],
                            st.session_state.interactive_gen_tokenizer,
                            st.session_state.interactive_gen_model
                        )

            if result.get("interview_questions"):
                with st.expander("🤔 AI-Generated Interview Questions", expanded=False):
                    st.markdown(result["interview_questions"])

            if result.get("highlighted_project"):
                with st.expander("🏆 Most Relevant Project Highlight", expanded=True):
                    project_data = result['highlighted_project']
                    st.success(f"Found most relevant project with a {project_data['score']:.2f}% semantic match.")
                    st.markdown(f"> {project_data['project']}")
        else:
            st.info("Your analysis results will appear here once you provide a resume and job description.")

    # --- Other Tabs (no changes needed) ---
    with tab2:
        st.header("🧠 Train Model")
        st.info(f"Model will be fine-tuned on '{CHECKPOINT}' and saved to '{FINE_TUNED_MODEL_PATH}'.")
        st.markdown("---")
        st.subheader("GPU Status")
        if torch.cuda.is_available():
            st.success(f"✅ CUDA available: {torch.cuda.get_device_name(0)}")
        else:
            st.error("❌ CUDA NOT available. Training will be slow.")
        st.markdown("---")
        raw_dataset = fallback_load_training_data()
        if st.button("🚀 Start Training", key="start_training_button"):
            if raw_dataset is None:
                st.error("❌ No training data loaded. Check 'datasets' folder.")
                return
            with st.spinner("Training model..."):
                tokenizer = AutoTokenizer.from_pretrained(CHECKPOINT)

                def tokenize(x): return tokenizer(x['text'], truncation=True, padding='max_length', max_length=512)

                tokenized_dataset = raw_dataset.map(tokenize, batched=True, remove_columns=["text"])
                training_model = AutoModelForMaskedLM.from_pretrained(CHECKPOINT)
                if torch.cuda.is_available(): torch.cuda.empty_cache()
                training_args = TrainingArguments(
                    output_dir=FINE_TUNED_MODEL_PATH, per_device_train_batch_size=8,
                    num_train_epochs=3, logging_dir='./logs', fp16=torch.cuda.is_available()
                )
                trainer = Trainer(
                    model=training_model, args=training_args, train_dataset=tokenized_dataset,
                    data_collator=DataCollatorForLanguageModeling(tokenizer=tokenizer, mlm_probability=0.15)
                )
                trainer.train()
                trainer.save_model(FINE_TUNED_MODEL_PATH)
            st.success("✅ Model training complete.")
    with tab3:
        st.header("📂 Batch Resume/JD Comparison")
        mode = st.radio("Select comparison mode:", [
            "📁 Resumes → 1 JD",
            "📁 1 Resume → JDs",
            "🧾 CSV of Pairs (text1/text2)"
        ], key="batch_mode_radio")

        uploaded_files = st.file_uploader("Upload resumes or JDs (PDF/DOCX/TXT)", type=["pdf", "docx", "txt"],
                                          accept_multiple_files=True, key="batch_files_upload")
        jd_input = st.text_area("Paste Job Description (for mode 1)", height=200, key="batch_jd_text")
        single_resume_input = st.text_area("Paste Resume (for mode 2)", height=200, key="batch_resume_text")
        pair_csv = st.file_uploader("Upload CSV with text1, text2 columns (for mode 3)", type=["csv"],
                                    key="batch_csv_upload")

        if st.button("⚙️ Run Batch Comparison", key="run_batch_button"):
            if (mode == "📁 Resumes → 1 JD" and (not jd_input or not uploaded_files)) or \
                    (mode == "📁 1 Resume → JDs" and (not single_resume_input or not uploaded_files)) or \
                    (mode == "🧾 CSV of Pairs (text1/text2)" and not pair_csv):
                st.warning("Please provide all necessary inputs for the selected batch mode.")
                st.stop()

            model = load_embedder(st.session_state.selected_model_path)
            results = []
            all_jd_keywords_for_heatmap = set()
            total_items = 0

            if mode != "🧾 CSV of Pairs (text1/text2)":
                total_items = len(uploaded_files)
            elif pair_csv:
                try:
                    df_temp = pd.read_csv(pair_csv)
                    total_items = df_temp.shape[0] if "text1" in df_temp.columns and "text2" in df_temp.columns else 0
                except (FileNotFoundError, pd.errors.ParserError) as e:
                    st.error(f"Invalid CSV file: {e}")
                    total_items = 0

            if total_items == 0:
                st.warning("No items to process in batch comparison.")
                st.stop()

            progress_bar = st.progress(0, text="Processing items...")
            processed_count = 0

            if mode == "📁 Resumes → 1 JD":
                jd_text_processed = jd_input
                active_jd_keyword_list_base = get_active_keywords(selected_keyword_system, user_defined_keywords)
                all_jd_keywords_for_heatmap.update(active_jd_keyword_list_base)

                for file in uploaded_files:
                    try:
                        resume_txt = textract.process(file.read()).decode('utf-8')
                        res = analyze(resume_txt, jd_input, model, user_keywords=user_defined_keywords,
                                      weights=(tfidf_weight, semantic_weight, keyword_weight),
                                      exclusions=exclusions, selected_keyword_system=selected_keyword_system)
                        res["Resume_File"] = file.name
                        results.append(res)
                    except Exception as e:
                        st.error(f"Error processing {file.name}: {e}")
                    processed_count += 1
                    progress_bar.progress(processed_count / total_items,
                                          text=f"Processing resumes... {processed_count}/{total_items}")

            elif mode == "📁 1 Resume → JDs":
                single_resume_text_processed = single_resume_input
                for file in uploaded_files:
                    try:
                        jd_txt = textract.process(file.read()).decode('utf-8')
                        general_jd_keywords_base = extract_keywords(clean_text(jd_txt.lower(), exclusions))
                        all_jd_keywords_for_heatmap.update(general_jd_keywords_base)

                        res = analyze(single_resume_text_processed, jd_txt, model, user_keywords=user_defined_keywords,
                                      weights=(tfidf_weight, semantic_weight, keyword_weight),
                                      exclusions=exclusions, selected_keyword_system=selected_keyword_system)
                        res["JD_File"] = file.name
                        results.append(res)
                    except Exception as e:
                        st.error(f"Error processing {file.name}: {e}")
                    processed_count += 1
                    progress_bar.progress(processed_count / total_items,
                                          text=f"Processing JDs... {processed_count}/{total_items}")

            elif mode == "🧾 CSV of Pairs (text1/text2)":
                df_pairs = pd.read_csv(pair_csv)
                if "text1" in df_pairs.columns and "text2" in df_pairs.columns:
                    for idx, row in df_pairs.iterrows():
                        try:
                            jd_txt_for_heatmap = str(row["text2"])
                            general_jd_keywords_base = extract_keywords(
                                clean_text(jd_txt_for_heatmap.lower(), exclusions))
                            all_jd_keywords_for_heatmap.update(general_jd_keywords_base)

                            res = analyze(str(row["text1"]), str(row["text2"]), model,
                                          user_keywords=user_defined_keywords,
                                          weights=(tfidf_weight, semantic_weight, keyword_weight),
                                          exclusions=exclusions, selected_keyword_system=selected_keyword_system)
                            res["Row_ID"] = idx + 1
                            results.append(res)
                        except Exception as e:
                            st.error(f"Error processing row {idx + 1}: {e}")
                        processed_count += 1
                        progress_bar.progress(processed_count / total_items,
                                              text=f"Processing CSV pairs... {processed_count}/{total_items}")
                else:
                    st.error("CSV must have 'text1' and 'text2' columns.")

            progress_bar.empty()

            if results:
                df_result = pd.DataFrame(results)
                df_display = df_result[
                    [col for col in df_result.columns if not isinstance(df_result[col].iloc[0], (dict, list))]]
                st.subheader("📊 Batch Comparison Results")
                st.dataframe(df_display)
                st.download_button("📥 Download Results as CSV", data=df_display.to_csv(index=False).encode('utf-8'),
                                   file_name="batch_comparison.csv", mime="text/csv", key="download_batch_csv")

                st.subheader("📈 Batch Insights")
                if not df_display.empty:
                    col1_viz, col2_viz = st.columns(2)
                    with col1_viz:
                        st.markdown("#### Final Score Distribution")
                        chart_score_dist = alt.Chart(df_display).mark_bar().encode(
                            x=alt.X('final_score:Q', bin=alt.Bin(maxbins=10), title="Final Score (%)"),
                            y=alt.Y('count()', title="Number of Matches")
                        ).properties(title='Distribution of Final Scores')
                        st.altair_chart(chart_score_dist, use_container_width=True)

                    with col2_viz:
                        st.markdown("#### Keyword Match Rate")
                        if all_jd_keywords_for_heatmap:
                            keyword_match_rates = []
                            for kw in all_jd_keywords_for_heatmap:
                                matched_count = sum(1 for r in results if kw in r['matched'])
                                match_rate = (matched_count / len(results)) * 100
                                keyword_match_rates.append({'Keyword': kw, 'Match Rate (%)': match_rate})

                            if keyword_match_rates:
                                heatmap_data = pd.DataFrame(keyword_match_rates).nlargest(20, 'Match Rate (%)')
                                chart_heatmap = alt.Chart(heatmap_data).mark_bar().encode(
                                    x=alt.X('Match Rate (%):Q'),
                                    y=alt.Y('Keyword:N', sort='-x'),
                                    tooltip=['Keyword', 'Match Rate (%)']
                                ).properties(title='Top 20 Keywords by Match Rate')
                                st.altair_chart(chart_heatmap, use_container_width=True)
    with tab4:
        st.header("🛠️ ATS Formatting Compatibility Check")
        st.info(
            "This tool scans your resume file for structural issues that can confuse older, less sophisticated Applicant Tracking Systems (ATS).")
        st.warning("Note: This check is for formatting only and does not analyze the text content.")

        uploaded_resume_for_check = st.file_uploader(
            "Upload your resume to check its formatting (DOCX or PDF recommended)",
            type=['docx', 'pdf'],
            key='format_check_upload'
        )

        if uploaded_resume_for_check:
            st.markdown("---")
            warnings = check_resume_formatting(uploaded_resume_for_check)

            if not warnings:
                st.success("✅ **Congratulations!** No common ATS formatting issues were found.")
                st.markdown("Your resume structure appears clean and should be easily parsable by most systems.")
            else:
                st.error(f"⚠️ **{len(warnings)} potential formatting issue(s) found:**")
                st.markdown(
                    "These issues may cause problems with older ATS. Consider creating a simplified, plain-text version of your resume for maximum compatibility.")
                for warning in warnings:
                    st.markdown(f"- {warning}")
    with tab5:
        st.header("📊 My Skills Inventory")
        st.info(
            "This dashboard scans your resume for technical skills defined in the keyword library and groups them by category.")

        if st.session_state.resume_text_content:
            resume_skills = extract_resume_skills(st.session_state.resume_text_content, KEYWORD_CATEGORIES)

            if not resume_skills:
                st.warning("No specific technical skills from the keyword library were found in your resume.")
            else:
                st.success(f"Found skills across {len(resume_skills)} categories.")
                col1, col2 = st.columns(2)
                sorted_categories = sorted(resume_skills.items())
                with col1:
                    for i in range(0, len(sorted_categories), 2):
                        category, skills = sorted_categories[i]
                        with st.expander(f"{category} ({len(skills)} found)", expanded=True):
                            for skill in sorted(skills):
                                st.markdown(f"- {skill.capitalize()}")
                with col2:
                    for i in range(1, len(sorted_categories), 2):
                        category, skills = sorted_categories[i]
                        with st.expander(f"{category} ({len(skills)} found)", expanded=True):
                            for skill in sorted(skills):
                                st.markdown(f"- {skill.capitalize()}")
        else:
            st.warning("Please upload or paste a resume in the 'Analyze Resume' tab to see your skills inventory.")
    with tab6:
        st.header("✉️ Enhanced Cover Letter Helper")
        st.info(
            "This tool helps you create or analyze a cover letter that is tailored to your resume and the job description.")

        if not st.session_state.analysis_result:
            st.warning("Please run an analysis on the 'Analyze Resume' tab first to enable the Cover Letter Helper.")
        else:
            result = st.session_state.analysis_result
            resume_text = st.session_state.resume_text_content
            jd_text = st.session_state.jd_text_content
            jd_keywords = result['missing'] + result['matched']
            mode = st.radio("Choose a mode:", ("✍️ Generate New Cover Letter", "🔍 Analyze Existing Cover Letter"),
                            horizontal=True)
            if mode == "✍️ Generate New Cover Letter":
                if st.button("🚀 Generate Draft", use_container_width=True):
                    with st.spinner("AI is drafting your cover letter..."):
                        gen_tokenizer, gen_model = load_generative_model()
                        draft = generate_cover_letter(resume_text, jd_text, result['matched'], gen_tokenizer, gen_model)
                        st.text_area("Generated Draft (you can edit and copy this):", value=draft, height=400)
            elif mode == "🔍 Analyze Existing Cover Letter":
                st.subheader("Paste Your Cover Letter Below:")
                cover_letter_input = st.text_area("Your cover letter text:", height=250, label_visibility="collapsed")
                if st.button("🔬 Analyze Cover Letter", use_container_width=True):
                    if cover_letter_input:
                        with st.spinner("Analyzing..."):
                            feedback_points = analyze_cover_letter(cover_letter_input, jd_keywords)
                            st.subheader("Analysis Feedback:")
                            for point in feedback_points:
                                st.markdown(point)
                    else:
                        st.warning("Please paste your cover letter text into the box to analyze it.")


if __name__ == '__main__':
    if not os.path.exists("domain_keywords.py"):
        with open("domain_keywords.py", "w") as f_out: f_out.write('...')
    main()